#!/usr/bin/env python3
"""
generate_lesson_plan.py - 专业级详案Word文档生成器 v3.0
深度整合人教版《教师教学用书》和昌明师友《教学设计与指导》
生成包含逐句话术、预设与生成的完整详案（非简案）
用法: python3 generate_lesson_plan.py <json_file> [output_path]
"""
import json, sys, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), val.get('space', '0'))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _add_para(doc, text, bold=False, size=12, indent=0, align=None, color=None, font_name='宋体',
              first_indent=0, line_spacing=1.5, space_after=0):
    """添加带格式的段落"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    p.paragraph_format.line_spacing = line_spacing
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    return p


def _add_rich_para(doc, segments, indent=0, line_spacing=1.5):
    """添加富文本段落 segments: [(text, bold, size, color, font), ...]"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.line_spacing = line_spacing
    for seg in segments:
        text = seg[0]
        bold = seg[1] if len(seg) > 1 else False
        size = seg[2] if len(seg) > 2 else 12
        color = seg[3] if len(seg) > 3 else None
        font = seg[4] if len(seg) > 4 else '宋体'
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = font
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font)
        if color:
            run.font.color.rgb = color
    return p


def _add_heading(doc, text, level=1):
    """添加带格式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading


def _add_separator(doc, color=RGBColor(200, 200, 200)):
    """添加分隔线"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = color


def _make_table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        cell.text = h
        _set_cell_shading(cell, 'D9E2F3')
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.cell(r_idx + 1, c_idx)
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    return t


def create_lesson_plan(data: dict, output_path: str):
    """生成专业级详案Word文档"""
    doc = Document()

    # === 页面设置 A4 ===
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    grade = data.get('grade', '___')
    semester = data.get('semester', '___')
    subject = data.get('subject', '___')
    edition = data.get('edition', '人教版')
    unit = data.get('unit', '___')
    lesson = data.get('lesson', '')
    topic = data.get('topic', '___')
    unit_name = data.get('unit_name', '')
    textbook_page = data.get('textbook_page', '')

    # ==================== 大标题 ====================
    title = doc.add_heading('', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'教    案（详案）')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    _add_para(doc, f'—— {grade}年级{semester} {subject} {edition} · {topic}',
              size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(100, 100, 100), font_name='楷体')
    _add_para(doc, f'整合人民教育出版社《教师教学用书》+ 昌明师友《教学设计与指导》详案内容',
              size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(150, 150, 150), font_name='楷体')

    # ==================== 基本信息表格 ====================
    doc.add_paragraph()
    info_data = [
        ('年    级', f'{grade}年级'),
        ('学    期', semester),
        ('学    科', subject),
        ('教材版本', edition),
        ('单    元', f'第{unit}单元' + (f'  {unit_name}' if unit_name else '')),
        ('课    题', topic),
        ('课    时', f'第{lesson}课时' if lesson else '共___课时'),
        ('教材页码', f'第{textbook_page}页' if textbook_page else '第___页'),
    ]

    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    idx = 0
    for r in range(4):
        for c in range(0, 4, 2):
            if idx < len(info_data):
                label_cell = info_table.cell(r, c)
                value_cell = info_table.cell(r, c + 1)
                _set_cell_shading(label_cell, 'F0F0F0')
                label_cell.text = info_data[idx][0]
                value_cell.text = info_data[idx][1]
                for cell in [label_cell, value_cell]:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(11)
                            run.font.name = '宋体'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                for run in label_cell.paragraphs[0].runs:
                    run.font.bold = True
                idx += 1

    doc.add_paragraph()

    # ==================== 一、教材分析 ====================
    _add_heading(doc, '一、教材分析', level=1)

    textbook_analysis = data.get('textbook_analysis', '')
    if textbook_analysis:
        _add_para(doc, textbook_analysis, first_indent=0.75, size=12)
    else:
        # 提供详案级结构化模板
        sections_data = [
            ('1. 本课在教材中的地位',
             '本课是___年级___册第___单元的第___课。本单元以"___"为主题，'
             '共编排了___篇课文/___个教学内容。本课在单元中处于___位置，'
             '前承___（前一课/前一单元的核心知识），为本课学习奠定了___基础；'
             '后启___（后续课/后续单元的知识），本课的学习将为___做好铺垫。\n\n'
             '【人教版教师教学用书要点】本课教材编写的主要意图在于___，'
             '体现了课程标准中"___"的要求。\n\n'
             '【昌明师友·教学设计要点】本课教学应重点关注___，'
             '建议采用___的方法突破重难点。'),
            ('2. 教材内容深度分析',
             '本课教材主要包含以下核心内容：\n'
             '（1）___——这是本课的基础知识/基本技能\n'
             '（2）___——这是本课的核心内容/重点所在\n'
             '（3）___——这是本课的拓展延伸/能力提升\n\n'
             '教材通过___的方式呈现，注重___能力的培养。'
             '教材中的插图/例题/活动设计旨在___。'),
            ('3. 编写意图与课标对接',
             '教材编者将本课安排在此处，意图在于：\n'
             '①___（知识建构层面）\n'
             '②___（能力发展层面）\n'
             '③___（素养培育层面）\n\n'
             '对应课程标准要求：___\n'
             '体现了___的课程理念。'),
            ('4. 知识体系图谱',
             '本课涉及的核心知识点及其逻辑关系：\n'
             '├── 核心知识点A：___\n'
             '│   ├── 子知识点A1：___\n'
             '│   └── 子知识点A2：___\n'
             '├── 核心知识点B：___\n'
             '│   ├── 子知识点B1：___\n'
             '│   └── 子知识点B2：___\n'
             '└── 核心知识点C：___\n\n'
             '知识点之间的逻辑关系为___，教学时应按照___的顺序推进。'),
        ]
        for sec_title, sec_content in sections_data:
            _add_rich_para(doc, [(sec_title, True, 13, RGBColor(0, 51, 102), '黑体')])
            _add_para(doc, sec_content, first_indent=0.75)
            doc.add_paragraph()

    # ==================== 二、学情分析 ====================
    _add_heading(doc, '二、学情分析', level=1)

    student_analysis = data.get('student_analysis', '')
    if student_analysis:
        _add_para(doc, student_analysis, first_indent=0.75)
    else:
        items = [
            ('1. 知识基础', '学生已经学习了___，掌握了___，具备了___的基础知识和基本技能。'
             '但对___的理解还不够深入，需要在本课中进一步巩固和提升。'),
            ('2. 认知特点', '___年级学生正处于___阶段（具体运算阶段/前运算阶段等），'
             '思维方式以___为主（形象思维为主，逐步向抽象思维过渡），'
             '对___较为敏感（直观、生动、有趣的内容），注意力持续时间约___分钟。'),
            ('3. 学习习惯', '学生已初步形成___的学习习惯（如：认真听讲、积极发言、合作学习等），'
             '但在___方面仍需加强（如：独立思考、自主探究、规范表达等）。'
             '班级整体学习氛围___，学优生约___人，学困生约___人。'),
            ('4. 困难预判', '基于对学生的了解，预判本课学习中学生可能遇到以下困难：\n'
             '   ①___（认知层面的困难）\n'
             '   ②___（技能层面的困难）\n'
             '   ③___（情感态度层面的困难）\n\n'
             '   针对以上困难，将在教学过程中通过___予以化解。'),
        ]
        for sec_title, sec_content in items:
            _add_rich_para(doc, [(sec_title + '：', True, 12)])
            _add_para(doc, sec_content, first_indent=0.75)
    doc.add_paragraph()

    # ==================== 三、教学目标 ====================
    _add_heading(doc, '三、教学目标', level=1)
    _add_para(doc, '【整合说明】以下目标综合了人教版《教师教学用书》的课程标准要求和昌明师友《教学设计与指导》的实践导向表述。',
              size=10, color=RGBColor(120, 120, 120), font_name='楷体')

    objectives = data.get('objectives', {})
    if objectives:
        dims = [
            ('（一）知识与技能目标', 'knowledge',
             '学生能够识记、理解并运用以下知识和技能：'),
            ('（二）过程与方法目标', 'process',
             '学生在学习过程中，通过___方法，培养以下能力：'),
            ('（三）情感态度与价值观目标', 'emotion',
             '通过本课学习，学生能够：'),
        ]
        for dim_title, dim_key, dim_desc in dims:
            _add_rich_para(doc, [(dim_title, True, 13, RGBColor(0, 51, 102), '黑体')])
            _add_para(doc, dim_desc, size=11, color=RGBColor(80, 80, 80))
            items = objectives.get(dim_key, [])
            if items:
                for i, item in enumerate(items, 1):
                    _add_para(doc, f'  {i}. {item}', size=11, indent=0.5)
            else:
                _add_para(doc, f'  （待根据教材分析确定）', size=11, indent=0.5, color=RGBColor(150, 150, 150))
    else:
        dims_template = [
            ('（一）知识与技能目标', [
                '1. 认识并会写本课生字词___个，理解由生字组成的词语和重点句子的意思。',
                '2. 正确、流利、有感情地朗读课文，背诵___段落。',
                '3. 理解课文主要内容，能用自己的话说说___。',
            ]),
            ('（二）过程与方法目标', [
                '1. 通过朗读品味语言，体会文章的表达方式（如___）。',
                '2. 通过讨论交流、合作探究，深化对文本的理解。',
                '3. 学习___的方法，培养___的能力。',
            ]),
            ('（三）情感态度与价值观目标', [
                '1. 感受___（文本中蕴含的情感/道理/美）。',
                '2. 培养___的良好品质/习惯/态度。',
                '3. 激发对___的兴趣和热爱。',
            ]),
        ]
        for dim_title, dim_items in dims_template:
            _add_rich_para(doc, [(dim_title, True, 13, RGBColor(0, 51, 102), '黑体')])
            for item in dim_items:
                _add_para(doc, item, size=11, indent=0.5)
    doc.add_paragraph()

    # ==================== 四、教学重点 ====================
    _add_heading(doc, '四、教学重点', level=1)
    key_points = data.get('key_points', [])
    key_strategies = data.get('key_point_strategies', [])
    if key_points:
        for i, kp in enumerate(key_points, 1):
            _add_rich_para(doc, [(f'重点{i}：', True, 12, RGBColor(180, 0, 0)), (kp, False, 12)])
            if i - 1 < len(key_strategies) and key_strategies[i - 1]:
                _add_rich_para(doc, [('  📌 突破策略：', True, 11, RGBColor(0, 100, 0)),
                                      (key_strategies[i - 1], False, 11)])
            doc.add_paragraph()
    else:
        _add_para(doc, '（待根据教材分析确定教学重点）')
    doc.add_paragraph()

    # ==================== 五、教学难点 ====================
    _add_heading(doc, '五、教学难点', level=1)
    difficulties = data.get('difficulties', [])
    diff_methods = data.get('difficulty_methods', [])
    if difficulties:
        for i, d in enumerate(difficulties, 1):
            _add_rich_para(doc, [(f'难点{i}：', True, 12, RGBColor(180, 0, 0)), (d, False, 12)])
            if i - 1 < len(diff_methods) and diff_methods[i - 1]:
                _add_rich_para(doc, [('  🔑 化解方法：', True, 11, RGBColor(0, 100, 0)),
                                      (diff_methods[i - 1], False, 11)])
            doc.add_paragraph()
    else:
        _add_para(doc, '（待根据教材分析确定教学难点）')
    doc.add_paragraph()

    # ==================== 六、教学方法 ====================
    _add_heading(doc, '六、教学方法', level=1)
    methods = data.get('teaching_methods', {})
    if methods:
        _add_rich_para(doc, [('（一）教法', True, 13, RGBColor(0, 51, 102), '黑体')])
        for m in methods.get('teaching', []):
            _add_para(doc, f'  • {m}', indent=0.5)
        _add_rich_para(doc, [('（二）学法', True, 13, RGBColor(0, 51, 102), '黑体')])
        for m in methods.get('learning', []):
            _add_para(doc, f'  • {m}', indent=0.5)
    else:
        _add_para(doc, '  教法：情境教学法、启发式教学法、直观演示法、任务驱动法')
        _add_para(doc, '  学法：自主探究、合作学习、实践操作、归纳总结')
    doc.add_paragraph()

    # ==================== 七、教学准备 ====================
    _add_heading(doc, '七、教学准备', level=1)
    preps = data.get('preparations', {})
    if isinstance(preps, dict):
        _add_rich_para(doc, [('教师准备：', True, 12)])
        for p in preps.get('teacher', ['多媒体课件', '板书预设', '相关教具']):
            _add_para(doc, f'  • {p}', style='List Bullet')
        _add_rich_para(doc, [('学生准备：', True, 12)])
        for p in preps.get('student', ['预习课文', '准备学具']):
            _add_para(doc, f'  • {p}', style='List Bullet')
    else:
        _add_para(doc, '  • 多媒体课件  • 相关教具  • 板书预设')
    doc.add_paragraph()

    # ==================== 八、教学过程（详案核心） ====================
    _add_heading(doc, '八、教学过程（详案）', level=1)
    _add_para(doc, '【说明】以下为完整详案，包含教师逐句话术、学生预判反应、预设与生成、课堂调控策略。可直接用于课堂执行。',
              size=10, color=RGBColor(0, 100, 0), font_name='楷体')

    process = data.get('teaching_process', [])
    if process:
        for step_idx, step in enumerate(process, 1):
            step_title = step.get('title', f'环节{step_idx}')
            step_time = step.get('time', '')
            step_content = step.get('content', '')
            teacher_act = step.get('teacher_activity', '')
            student_act = step.get('student_activity', '')
            design_intent = step.get('design_intent', '')
            preset_gen = step.get('preset_generation', '')
            transition = step.get('transition', '')
            classroom_mgmt = step.get('classroom_management', '')

            # 环节标题（醒目）
            _add_separator(doc, RGBColor(0, 51, 102))
            _add_rich_para(doc, [
                (f'环节{step_idx}：{step_title}', True, 16, RGBColor(0, 51, 102), '黑体'),
                (f'  ⏱ {step_time}', False, 12, RGBColor(100, 100, 100)),
            ])

            if step_content:
                _add_rich_para(doc, [('【教学内容】', True, 11, RGBColor(80, 80, 80), '黑体')])
                _add_para(doc, step_content, size=11, first_indent=0.5)

            # 师生活动表格（详案级）
            if teacher_act or student_act:
                t = doc.add_table(rows=2, cols=2)
                t.style = 'Table Grid'
                t.alignment = WD_TABLE_ALIGNMENT.CENTER

                _set_cell_shading(t.cell(0, 0), 'D9E2F3')
                _set_cell_shading(t.cell(0, 1), 'E8F5E9')
                t.cell(0, 0).text = '📋 教师活动（逐句话术）'
                t.cell(0, 1).text = '👦👧 学生活动（预判反应）'
                t.cell(1, 0).text = teacher_act if teacher_act else '（待填写）'
                t.cell(1, 1).text = student_act if student_act else '（待填写）'

                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(10.5)
                                run.font.name = '宋体'
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                for run in t.cell(0, 0).paragraphs[0].runs:
                    run.font.bold = True
                for run in t.cell(0, 1).paragraphs[0].runs:
                    run.font.bold = True

            # 设计意图
            if design_intent:
                _add_rich_para(doc, [('【设计意图】', True, 11, RGBColor(0, 128, 0), '黑体')])
                _add_para(doc, design_intent, size=11, indent=0.3, first_indent=0.5)

            # 预设与生成（详案核心）
            if preset_gen:
                _add_rich_para(doc, [('【预设与生成】', True, 11, RGBColor(200, 100, 0), '黑体')])
                _add_para(doc, preset_gen, size=11, indent=0.3, first_indent=0.5)

            # 过渡语
            if transition:
                _add_rich_para(doc, [('【过渡语】', True, 11, RGBColor(100, 0, 150), '黑体'),
                                      (f' "{transition}"', False, 11, RGBColor(100, 0, 150))])

            # 课堂调控
            if classroom_mgmt:
                _add_rich_para(doc, [('【课堂调控】', True, 11, RGBColor(0, 100, 150), '黑体')])
                _add_para(doc, classroom_mgmt, size=10, indent=0.3)

            doc.add_paragraph()  # 环节间空行
    else:
        # 默认6个环节的详案模板
        default_steps = [
            {
                'title': '情境导入',
                'time': '5分钟',
                'content': '通过创设与课题相关的生活情境，激发学生学习兴趣，建立新旧知识的联系。',
                'teacher': (
                    '① 教师话术："同学们，老师今天给大家带来了一样神奇的东西（出示/展示___）。'
                    '你们知道这是什么吗？"\n'
                    '② 追问："有谁见过/听说过___？能给大家说说吗？"\n'
                    '③ 引导："看来同学们对___很感兴趣。那你们想不想知道___？"\n'
                    '④ 过渡语："好，今天我们就一起来学习——《___》。"（板书课题）\n'
                    '⑤ 板书动作：在黑板正中书写课题，字迹工整，大小适中。'
                ),
                'student': (
                    '① 预设A：学生积极举手回答 → 教师："你观察得真仔细！还有谁想补充？"\n'
                    '② 预设B：学生沉默 → 教师提示："想想看，你在___的时候有没有见过___？"\n'
                    '③ 预设C：学生回答偏离 → 教师引导："你说得很有趣，不过我们现在先来关注___。"'
                ),
                'intent': '从学生已有的生活经验出发，创设真实情境，激发学习兴趣和探究欲望。'
                         '符合课标"联系生活实际，创设真实情境"的教学建议。',
                'preset': (
                    '预设A（顺利）：学生能说出___→ 教师跟进："说得很好！那你们知道它是怎么来的吗？"\n'
                    '预设B（部分）：学生只能说出表面 → 教师追问："除了___，你还发现了什么？"\n'
                    '预设C（困难）：学生完全不了解 → 教师降低难度："没关系，我们先来看看___。"'
                ),
                'transition': '刚才同学们都观察得很仔细，接下来，让我们打开课本，一起走进___的世界。',
                'mgmt': '导入阶段控制在5分钟以内，避免拖沓。关注后排学生的参与度，适当走动巡视。'
            },
            {
                'title': '新知探究',
                'time': '18分钟',
                'content': '教师引导学生逐步探究本课核心知识，通过"感知→理解→掌握"的认知规律，让学生在亲身经历中建构知识。',
                'teacher': (
                    '① 任务布置："现在请大家打开课本第___页，自由朗读课文，注意读准字音，读通句子。"\n'
                    '② 巡视指导：关注学生的朗读情况，个别指导读音困难的学生。\n'
                    '③ 检查反馈："好，我们来看看这些生字词。（出示PPT）谁来读一读？"\n'
                    '④ 精读引导："课文第___段写了___，请同学们仔细读一读，'
                    '找一找哪些词句最能体现___？用笔画出来。"\n'
                    '⑤ 点拨提升："同学们找得很好。我们来看___这个词/句，'
                    '它用了___的方法，写出了___。"\n'
                    '⑥ 板书核心知识：在黑板上结构化呈现本课知识要点。'
                ),
                'student': (
                    '① 学生自由朗读课文，标注生字词和重点句子。\n'
                    '② 预设A：学生准确找出关键词句 → 教师追问"为什么这个词/句能体现___？"\n'
                    '③ 预设B：学生找得不够准确 → 教师提示"再读一遍第___句，想想___"\n'
                    '④ 预设C：学生提出意想不到的好问题 → 教师肯定并全班讨论\n'
                    '⑤ 学生记录笔记，标注老师强调的重点。'
                ),
                'intent': '遵循"感知→理解→掌握"的认知规律。先让学生充分接触文本（感知），'
                         '再通过师生互动深入分析（理解），最后归纳总结形成知识体系（掌握）。',
                'preset': (
                    '预设A（读音问题）：学生读错___字 → 教师："这个字的读音是___，注意___的发音。"\n'
                    '预设B（理解偏差）：学生理解为___ → 教师引导："你的理解有道理，但我们再看看上下文，'
                    '作者想表达的是___。"\n'
                    '预设C（深度提问）：学生问"___？" → 教师赞赏："这个问题提得很好！'
                    '谁能帮他解答？"或教师直接解答。'
                ),
                'transition': '通过刚才的学习，我们了解了___。接下来，请同学们小组合作，讨论___。',
                'mgmt': '新知探究环节是本课核心，要保证充足时间（18分钟左右）。'
                       '朗读环节可采用多种形式：自由读、指名读、齐读交替进行。'
                       '精读环节注意给学生足够的思考时间（提问后等待3-5秒再请人回答）。'
            },
            {
                'title': '合作交流',
                'time': '7分钟',
                'content': '组织小组合作学习，在交流中深化理解，培养合作能力和表达能力。',
                'teacher': (
                    '① 分组："请前后桌4人为一小组，讨论___问题。"'
                    '（PPT出示讨论问题和要求）\n'
                    '② 要求说明："讨论时请注意：一人记录，一人汇报。'
                    '讨论时间4分钟，开始！"\n'
                    '③ 巡视指导：深入各组倾听讨论，必要时参与讨论引导方向。\n'
                    '④ 汇报组织："时间到！哪个小组愿意先来汇报？"\n'
                    '⑤ 点评提升：（针对汇报内容）"你们组的发现很有价值！'
                    '特别是___这一点，说得非常好。其他组有补充吗？"'
                ),
                'student': (
                    '① 小组内分工讨论，记录员记录讨论结果。\n'
                    '② 预设A：讨论热烈，结论多样 → 教师对比各组结论，引导发现差异\n'
                    '③ 预设B：讨论冷场 → 教师降低问题难度或提供讨论支架\n'
                    '④ 预设C：讨论偏离主题 → 教师巡视时及时引导\n'
                    '⑤ 汇报代表展示讨论成果，其他组认真聆听并补充。'
                ),
                'intent': '通过合作学习培养团队协作能力，在交流中完善思维，实现思维的碰撞与提升。'
                         '同时培养学生的倾听能力和表达能力。',
                'preset': (
                    '预设A（各组结论一致）：教师肯定后追问"还有没有不同的看法？"\n'
                    '预设B（各组结论不同）：教师引导对比分析"为什么会有不同的结论？'
                    '哪种更合理？为什么？"\n'
                    '预设C（讨论不到位）：教师提供思考支架"我们来从___角度想想看。"'
                ),
                'transition': '同学们讨论得很热烈，汇报也很精彩！通过合作学习，我们发现___。'
                            '现在，让我们来练一练，看看大家掌握得怎么样。',
                'mgmt': '讨论时间严格控制在4分钟，用计时器提醒。'
                       '巡视时重点关注学困生所在小组，给予适当帮助。'
                       '汇报环节注意控制时间，选择有代表性的2-3组即可。'
            },
            {
                'title': '巩固练习',
                'time': '8分钟',
                'content': '通过分层练习巩固所学知识，及时反馈，查漏补缺。',
                'teacher': (
                    '① 出示基础练习（PPT或学案）："请同学们独立完成练习第1-3题，时间3分钟。"\n'
                    '② 巡视收集典型错误，拍照或记录。\n'
                    '③ 集体讲评："我们一起来看第___题，大部分同学都做对了。'
                    '第___题有同学出错了，我们来看看错在哪里。"\n'
                    '④ 错误分析："这位同学的答案是___，你能说说你的想法吗？'
                    '（倾听后）原来你是这样想的。正确的方法应该是___。"\n'
                    '⑤ 提高练习："基础题完成得不错！现在挑战一下提高题。"'
                ),
                'student': (
                    '① 独立完成基础练习题。\n'
                    '② 预设A：全部正确 → 教师鼓励并引导挑战提高题\n'
                    '③ 预设B：部分错误 → 教师利用错误进行教学\n'
                    '④ 预设C：错误较多 → 教师重新讲解核心知识点\n'
                    '⑤ 学生自查订正，记录错题和正确方法。'
                ),
                'intent': '遵循"由易到难"的练习梯度设计。基础题确保全体学生掌握核心知识，'
                         '提高题为中等及以上学生提供思维挑战。及时反馈帮助学生查漏补缺。',
                'preset': (
                    '预设A（正确率高）：教师肯定后快速推进，增加提高题时间。\n'
                    '预设B（正确率一般）：教师重点讲解错误集中的题目。\n'
                    '预设C（错误率高）：教师暂停练习，回归知识点重新讲解后再练。'
                ),
                'transition': '通过练习，大家对___掌握得越来越好了。接下来，让我们一起回顾一下今天学到了什么。',
                'mgmt': '基础练习时间严格控制在3分钟内，避免超时。'
                       '讲评时以鼓励为主，不要在错误上过度纠缠。'
                       '关注学困生完成情况，必要时给予个别指导。'
            },
            {
                'title': '课堂总结',
                'time': '3分钟',
                'content': '引导学生梳理本课知识体系，形成完整认知。',
                'teacher': (
                    '① 引导回顾："同学们，这节课我们学习了《___》。谁能说说，你学到了什么？"\n'
                    '② 师生共同梳理：（指着板书）"我们一起回顾一下：'
                    '首先我们___，然后___，最后___。"\n'
                    '③ 升华提升："通过今天的学习，我们不仅知道了___，'
                    '更重要的是学会了___的方法。"\n'
                    '④ 情感引导："___（联系课文情感/价值观升华）。"'
                ),
                'student': (
                    '① 预设A：学生总结全面 → 教师补充完善\n'
                    '② 预设B：学生总结片面 → 教师追问"还有补充吗？"\n'
                    '③ 预设C：学生不发言 → 教师降低难度"我们来看看板书，'
                    '第一点是什么？"\n'
                    '④ 学生口述收获，完善笔记。'
                ),
                'intent': '通过系统梳理帮助学生形成完整的知识体系，加深理解和记忆。'
                         '总结环节让学生主动回忆，比教师直接告知更有利于知识内化。',
                'preset': '预设：学生总结可能不够全面 → 追问"还有补充吗？"引导查漏补缺。',
                'transition': '同学们总结得很好！今天的课就上到这里。',
                'mgmt': '总结环节不宜过长（3分钟），以学生口述为主。'
                       '板书此时起到关键的可视化梳理作用。'
            },
            {
                'title': '布置作业',
                'time': '2分钟',
                'content': '分层布置课后作业，满足不同层次学生的需求。',
                'teacher': (
                    '① 说明作业："今天的作业有三项，请大家记下来。"\n'
                    '② 基础作业（必做）："第一项，___。这是每个人都要完成的。"\n'
                    '③ 提高作业（选做）："第二项，___。有兴趣的同学可以挑战一下。"\n'
                    '④ 拓展作业（挑战）："第三项，___。学有余力的同学可以试试。"\n'
                    '⑤ 温馨提示："作业请在明天上课前完成。有问题随时问老师。"'
                ),
                'student': '学生记录作业内容，明确各项要求。',
                'intent': '分层作业设计满足不同水平学生的需求：基础作业巩固核心知识，'
                         '提高作业培养迁移能力，拓展作业激发探究欲望。',
                'preset': '',
                'transition': '',
                'mgmt': '布置作业要清晰、简洁，确保每个学生都记录了必做作业。'
                       '选做和挑战作业要强调"鼓励但不强制"。'
            },
        ]

        for step in default_steps:
            _add_separator(doc, RGBColor(0, 51, 102))
            _add_rich_para(doc, [
                (f'环节：{step["title"]}', True, 16, RGBColor(0, 51, 102), '黑体'),
                (f'  ⏱ {step["time"]}', False, 12, RGBColor(100, 100, 100)),
            ])

            if step['content']:
                _add_rich_para(doc, [('【教学内容】', True, 11, RGBColor(80, 80, 80), '黑体')])
                _add_para(doc, step['content'], size=11, first_indent=0.5)

            # 师生活动表格
            t = doc.add_table(rows=2, cols=2)
            t.style = 'Table Grid'
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_cell_shading(t.cell(0, 0), 'D9E2F3')
            _set_cell_shading(t.cell(0, 1), 'E8F5E9')
            t.cell(0, 0).text = '📋 教师活动（逐句话术）'
            t.cell(0, 1).text = '👦👧 学生活动（预判反应）'
            t.cell(1, 0).text = step['teacher']
            t.cell(1, 1).text = step['student']
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(10.5)
                            run.font.name = '宋体'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            for run in t.cell(0, 0).paragraphs[0].runs:
                run.font.bold = True
            for run in t.cell(0, 1).paragraphs[0].runs:
                run.font.bold = True

            if step['intent']:
                _add_rich_para(doc, [('【设计意图】', True, 11, RGBColor(0, 128, 0), '黑体')])
                _add_para(doc, step['intent'], size=11, first_indent=0.5)

            if step['preset']:
                _add_rich_para(doc, [('【预设与生成】', True, 11, RGBColor(200, 100, 0), '黑体')])
                _add_para(doc, step['preset'], size=11, first_indent=0.5)

            if step['transition']:
                _add_rich_para(doc, [('【过渡语】', True, 11, RGBColor(100, 0, 150), '黑体'),
                                      (f' "{step["transition"]}"', False, 11, RGBColor(100, 0, 150))])

            if step['mgmt']:
                _add_rich_para(doc, [('【课堂调控】', True, 11, RGBColor(0, 100, 150), '黑体')])
                _add_para(doc, step['mgmt'], size=10, indent=0.3)

            doc.add_paragraph()

    # ==================== 九、板书设计 ====================
    _add_heading(doc, '九、板书设计', level=1)
    _add_para(doc, '【说明】板书体现知识体系与逻辑关系，是课堂总结的重要视觉支架。',
              size=10, color=RGBColor(120, 120, 120), font_name='楷体')

    board = data.get('board_design', '')
    if board:
        p = doc.add_paragraph()
        run = p.add_run(board)
        run.font.size = Pt(11)
        run.font.name = 'Courier New'
        p.paragraph_format.left_indent = Cm(1)
    else:
        _add_para(doc, '┌─────────────────────────────────────┐', font_name='Courier New', indent=1)
        _add_para(doc, '│              《课题名》              │', font_name='Courier New', indent=1)
        _add_para(doc, '├──────────┬──────────┬──────────┤', font_name='Courier New', indent=1)
        _add_para(doc, '│  知识点1  │  知识点2  │  知识点3  │', font_name='Courier New', indent=1)
        _add_para(doc, '│  ······  │  ······  │  ······  │', font_name='Courier New', indent=1)
        _add_para(doc, '│  ······  │  ······  │  ······  │', font_name='Courier New', indent=1)
        _add_para(doc, '├──────────┴──────────┴──────────┤', font_name='Courier New', indent=1)
        _add_para(doc, '│         核心结论/方法总结         │', font_name='Courier New', indent=1)
        _add_para(doc, '└─────────────────────────────────────┘', font_name='Courier New', indent=1)
    doc.add_paragraph()

    # ==================== 十、作业设计 ====================
    _add_heading(doc, '十、作业设计', level=1)
    homework = data.get('homework', {})
    if homework:
        for level, items in homework.items():
            _add_rich_para(doc, [(level, True, 12)])
            if isinstance(items, list):
                for item in items:
                    _add_para(doc, f'  • {item}', indent=0.5)
            else:
                _add_para(doc, f'  • {items}', indent=0.5)
    else:
        levels = [
            ('★ 基础作业（必做）', [
                '抄写/背诵本课核心内容',
                '完成课后练习第___题',
            ]),
            ('★★ 提高作业（选做）', [
                '完成配套练习册相关习题',
                '用本课所学解释生活中的___现象',
            ]),
            ('★★★ 拓展作业（挑战）', [
                '查阅相关资料，了解___',
                '制作___小报/完成___实践任务',
            ]),
        ]
        for level, items in levels:
            _add_rich_para(doc, [(level, True, 12)])
            for item in items:
                _add_para(doc, f'  • {item}', indent=0.5)
    doc.add_paragraph()

    # ==================== 十一、教学反思 ====================
    _add_heading(doc, '十一、教学反思', level=1)
    reflection = data.get('reflection', {})
    if reflection:
        for key, val in reflection.items():
            _add_rich_para(doc, [(key, True, 12)])
            if val:
                _add_para(doc, val)
            else:
                for _ in range(4):
                    _add_para(doc, '_' * 55)
    else:
        reflection_items = [
            ('1. 预设成功点：', '本课设计中预计最有效的环节是___，因为___'),
            ('2. 实际亮点：', '（课后填写课堂中的精彩瞬间和学生表现）'),
            ('3. 需改进处：', '（课后填写需要调整的教学环节和方法）'),
            ('4. 学生反馈：', '（课后填写学生的学习效果和感受）'),
            ('5. 改进措施：', '（课后填写下一课/下一循环的具体改进方案）'),
        ]
        for title, content in reflection_items:
            _add_rich_para(doc, [(title, True, 12)])
            if content:
                _add_para(doc, content)
            else:
                for _ in range(3):
                    _add_para(doc, '_' * 55)

    # ==================== 资料来源说明 ====================
    doc.add_page_break()
    _add_heading(doc, '附：资料来源说明', level=1)
    sources = data.get('sources_used', [
        '1. 人民教育出版社《教师教学用书》（电子版）—— 教材分析、教学目标、教学建议',
        '2. 昌明师友·华东师范大学出版社《教学设计与指导》（电子版）—— 完整教学设计、师生活动详案',
        '3. 教材原文及相关教参',
        '4. 课程标准（2022年版）',
    ])
    for s in sources:
        _add_para(doc, s)
    doc.add_paragraph()
    _add_para(doc, '本教案由备课系统自动生成，整合两大核心教参的详案内容，仅供参考，请根据实际教学情况调整使用。',
              size=10, color=RGBColor(120, 120, 120), font_name='楷体')

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python3 generate_lesson_plan.py <json_file> [output_path]")
        sys.exit(1)
    with open(sys.argv[1], 'r', encoding='utf-8') as f:
        data = json.load(f)
    output = sys.argv[2] if len(sys.argv) > 2 else sys.argv[1].replace('.json', '_教案.docx')
    print(f"教案已生成: {create_lesson_plan(data, output)}")
